"""Build the answers report as a Word document and a PDF from report_content.BLOCKS.

Questions: Times New Roman 14, bold.  Answers: Times New Roman 12, regular.
"""
import pathlib
from PIL import Image

from report_content import BLOCKS, TITLE, SUBTITLE

ROOT = pathlib.Path(__file__).parent
FIGS = ROOT / "partC_output"
DOCX = ROOT / "CS3621_L05_Answers.docx"
PDF = ROOT / "CS3621_L05_Answers.pdf"

TNR = "/System/Library/Fonts/Supplemental/Times New Roman.ttf"
TNR_B = "/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf"
TNR_I = "/System/Library/Fonts/Supplemental/Times New Roman Italic.ttf"


# --------------------------------------------------------------------- Word
def build_docx():
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(1)
        sec.left_margin = sec.right_margin = Inches(1)
    usable = doc.sections[0].page_width.inches - 2

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    def para(text, size=12, bold=False, italic=False, space_before=0, align=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        if align is not None:
            p.alignment = align
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = RGBColor(0, 0, 0)
        return p

    para(TITLE, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    if SUBTITLE:
        para(SUBTITLE, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

    for block in BLOCKS:
        kind = block[0]
        if kind == "part":
            para(block[1], size=14, bold=True, space_before=16)
        elif kind == "q":
            para(block[1], size=14, bold=True, space_before=14)
        elif kind == "a":
            para(block[1], size=12)
        elif kind == "img":
            path = FIGS / block[1]
            w, h = Image.open(path).size
            width = min(usable, 6.3)
            doc.add_picture(str(path), width=Inches(width))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            para(block[2], size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        elif kind == "table":
            rows = block[1]
            t = doc.add_table(rows=len(rows), cols=len(rows[0]))
            t.style = "Table Grid"
            for i, row in enumerate(rows):
                for j, cell in enumerate(row):
                    c = t.cell(i, j)
                    c.text = ""
                    p = c.paragraphs[0]
                    p.paragraph_format.space_after = Pt(2)
                    r = p.add_run(str(cell))
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(10)
                    r.bold = (i == 0)
                    r.font.color.rgb = RGBColor(0, 0, 0)
            doc.add_paragraph()

    doc.save(DOCX)
    print("wrote", DOCX)


# ---------------------------------------------------------------------- PDF
def build_pdf():
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Image as RLImage, Table, TableStyle, KeepTogether)

    pdfmetrics.registerFont(TTFont("TNR", TNR))
    pdfmetrics.registerFont(TTFont("TNR-Bold", TNR_B))
    pdfmetrics.registerFont(TTFont("TNR-Italic", TNR_I))
    pdfmetrics.registerFontFamily("TNR", normal="TNR", bold="TNR-Bold", italic="TNR-Italic")

    margin = 2.0 * cm
    frame_w = A4[0] - 2 * margin

    q_style = ParagraphStyle("q", fontName="TNR-Bold", fontSize=14, leading=17,
                             spaceBefore=14, spaceAfter=6, textColor=colors.black)
    a_style = ParagraphStyle("a", fontName="TNR", fontSize=12, leading=15,
                             spaceAfter=6, textColor=colors.black)
    cap_style = ParagraphStyle("cap", fontName="TNR-Italic", fontSize=10, leading=12,
                               alignment=TA_CENTER, spaceAfter=8, textColor=colors.black)
    cell_style = ParagraphStyle("cell", fontName="TNR", fontSize=10, leading=12,
                                textColor=colors.black)
    head_style = ParagraphStyle("head", fontName="TNR-Bold", fontSize=10, leading=12,
                                textColor=colors.black)
    title_style = ParagraphStyle("t", fontName="TNR-Bold", fontSize=14, leading=17,
                                 alignment=TA_CENTER, spaceAfter=4, textColor=colors.black)
    sub_style = ParagraphStyle("s", fontName="TNR", fontSize=12, leading=15,
                               alignment=TA_CENTER, spaceAfter=14, textColor=colors.black)
    sub_title_only = ParagraphStyle("t2", parent=title_style, spaceAfter=14)

    def esc(t):
        return (str(t).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))

    story = [Paragraph(esc(TITLE), title_style if SUBTITLE else sub_title_only)]
    if SUBTITLE:
        story.append(Paragraph(esc(SUBTITLE), sub_style))

    for block in BLOCKS:
        kind = block[0]
        if kind in ("part", "q"):
            story.append(Paragraph(esc(block[1]).replace("\n", "<br/>"), q_style))
        elif kind == "a":
            story.append(Paragraph(esc(block[1]), a_style))
        elif kind == "img":
            path = FIGS / block[1]
            w, h = Image.open(path).size
            width = frame_w
            height = width * h / w
            max_h = A4[1] - 2 * margin - 60
            if height > max_h:
                height = max_h
                width = height * w / h
            story.append(KeepTogether([
                RLImage(str(path), width=width, height=height),
                Spacer(1, 4),
                Paragraph(esc(block[2]), cap_style),
            ]))
        elif kind == "table":
            rows = block[1]
            ncol = len(rows[0])
            data = [[Paragraph(esc(c), head_style if i == 0 else cell_style)
                     for c in row] for i, row in enumerate(rows)]
            # every column must at least fit its longest single word unbroken;
            # the slack left over is shared out by how much text each column carries
            pad = 9
            mins, loads = [], []
            for j in range(ncol):
                widest_word = 0.0
                for i, r in enumerate(rows):
                    font = "TNR-Bold" if i == 0 else "TNR"
                    for word in str(r[j]).split():
                        widest_word = max(widest_word, pdfmetrics.stringWidth(word, font, 10))
                mins.append(widest_word + pad)
                loads.append(sum(len(str(r[j])) for r in rows))
            slack = frame_w - sum(mins)
            if slack > 0:
                total_load = sum(loads) or 1
                widths = [m + slack * l / total_load for m, l in zip(mins, loads)]
            else:  # too wide to honour every minimum: fall back to proportional
                scale = frame_w / sum(mins)
                widths = [m * scale for m in mins]
            t = Table(data, colWidths=widths, repeatRows=1)
            t.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.4, colors.black),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]))
            story.append(KeepTogether([t]) if len(rows) <= 8 else t)
            story.append(Spacer(1, 10))

    def footer(canvas, doc_):
        canvas.saveState()
        canvas.setFont("TNR", 10)
        canvas.setFillColor(colors.black)
        canvas.drawCentredString(A4[0] / 2, margin / 2, str(canvas.getPageNumber()))
        canvas.restoreState()

    SimpleDocTemplate(str(PDF), pagesize=A4, topMargin=margin, bottomMargin=margin,
                      leftMargin=margin, rightMargin=margin,
                      title="CS3621 L05 Practical Answers").build(story, onFirstPage=footer,
                                                                  onLaterPages=footer)
    print("wrote", PDF)


if __name__ == "__main__":
    build_docx()
    build_pdf()
